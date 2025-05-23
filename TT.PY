import streamlit as st
import openai
import random
import datetime
import os
import json
from docx import Document

# Configuration de ta clé OpenAI
openai.api_key = "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"
# 📁 Chemin du fichier DOCX du sujet complet
chemin_sujet = "C:/Users/ps3ka/openai/BLOC1 DS.docx"

# 📂 Chargement des profils de fautes
with open("profils_fautes_corriges.json", "r", encoding="utf-8") as f:
    profils_fautes = json.load(f)
# 🔠 Lecture du contenu texte du fichier DOCX
def lire_docx_texte(chemin):
    doc = Document(chemin)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return "\n".join(full_text)

# ✍️ Génération du prompt académique à partir de la partie choisie
def construire_prompt_complet(sujet_complet, partie_cible, genre):
    intro = (
        "Tu incarnes un(e) étudiant(e) en Master 2 Finance de Marché – spécialité Trading. "
        "rédigeant un devoir surveillé basé sur un sujet fourni, en adoptant un ton personnel et critique, "
        "à la première personne du singulier, mais avec la rigueur attendue à ce niveau."
        "Tu dois démontrer une réflexion critique, une rigueur académique, et une bonne maîtrise des concepts financiers."
    )
    ton = (
        "Exprime-toi comme une étudiante en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiante…'. "
        "Sois précise, nuancée, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
        if genre == "f" else
        "Exprime-toi comme un étudiant en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiant…'. "
        "Sois précis, nuancé, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
    )
    consignes = (
        f"La partie {partie_cible} comprend plusieurs tâches (par exemple : veille économique, veille juridique, etc.). "
        "Tu dois impérativement répondre à **chaque tâche séparément** avec une **réponse détaillée, critique et construite**, "
        
        "Structure chaque réponse comme une mini-dissertation avec une **introduction, un développement et une conclusion**. "
        "Utilise les données du sujet (macroéconomie, inflation, volatilité, contexte géopolitique, etc.) et exploite-les avec précision. "
        "Tu dois démontrer ta capacité à faire de la **veille stratégique**, à **croiser les données** et à **proposer une analyse pertinente**."
    )

    output = (
        f"{intro}\n{ton}\n{consignes}\n\n"
        f"=== Sujet Complet ===\n{sujet_complet}\n\n"
        f"=== Consigne spécifique ===\n"
        f"Tu dois rédiger uniquement la partie intitulée : {partie_cible}. "
        f"Traite chaque tâche individuellement, comme si tu rédigeais un mémoire académique."
    )

    return output
# 🧠 Ajout de fautes personnalisées
def introduire_erreurs_subtiles(texte, profil):
    base = profil["base_rate"]
    mots_risque = profil["mots_a_risque"]
    repetition = profil["repetition_faute"]

    if random.random() < base["accent_oublie"]:
        texte = texte.replace("é", "e", 1).replace("à", "a", 1)
    if random.random() < base["homophone"]:
        homophones = [("et", "est"), ("a", "à"), ("ces", "ses"), ("on", "ont"), ("c’est", "s’est")]
        h = random.choice(homophones)
        texte = texte.replace(h[0], h[1], 1)
    if random.random() < base["touche_adj"]:
        texte = texte.replace("u", "i", 1)
    if random.random() < base["liaison_oubliee"]:
        texte = texte.replace("n’est", "ne est", 1)
    if random.random() < base["fautes_d_accord"]:
        texte = texte.replace("les entreprises sont", "les entreprise est", 1)
    if random.random() < repetition:
        mot = random.choice(mots_risque)
        texte = texte.replace(mot, mot.upper(), 1)

    return texte

# 🌐 Interface utilisateur Streamlit
st.title("🧠 Générateur de Devoirs Automatisé - Bloc 1")
st.markdown("Ce générateur traite un sujet complet et génère une réponse structurée pour une partie choisie.")

nom = st.text_input("Nom de l'élève")
prenom = st.text_input("Prénom de l'élève")
date_naissance = st.date_input("Date de naissance")
genre = st.radio("Genre", ["m", "f"])

noms_profils = [profil["nom"] for profil in profils_fautes.values()]
profil_nom_selectionne = st.selectbox("Choisis un profil d'élève :", noms_profils)
profil_selectionne = next(p for p in profils_fautes.values() if p["nom"] == profil_nom_selectionne)

partie_cible = st.selectbox("Choisis la partie à traiter :", [
    "Partie 1", "Partie 2", "Partie 3", "Partie 4", "Partie 5"
])

if st.button("📄 Générer le devoir"):
    try:
        texte_sujet = lire_docx_texte(chemin_sujet)
        prompt = construire_prompt_complet(texte_sujet, partie_cible, genre)
        temperature = random.uniform(0.7, 0.9)

        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=temperature,
            messages=[
                {"role": "system", "content": "Tu es professeur et expert en finance de marché. Tu évalues un devoir d’un étudiant en M2."},
                {"role": "user", "content": prompt}
            ]
        )

        texte_genere = response.choices[0].message.content.strip()
        texte_genere = introduire_erreurs_subtiles(texte_genere, profil_selectionne)

        st.subheader("✍️ Devoir généré")
        st.write(texte_genere)

        # 💾 Génération du fichier Word
        date_realisation = datetime.datetime.today().strftime("%d/%m/%Y")
        repertoire_eleve = os.path.join("C:/Users/ps3ka/openai/Devoirs_Eleves", f"{prenom}_{nom}")
        os.makedirs(repertoire_eleve, exist_ok=True)
        filename = os.path.join(repertoire_eleve, f"Devoir_{partie_cible}_{prenom}_{nom}.docx")

        doc = Document()
        doc.add_heading(f"Devoir - {partie_cible}", level=1)
        doc.add_paragraph(f"📌 Nom : {nom}")
        doc.add_paragraph(f"📌 Prénom : {prenom}")
        doc.add_paragraph(f"📌 Date de naissance : {date_naissance.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"📌 Date de réalisation : {date_realisation}")
        doc.add_heading("Réponse :", level=2)
        doc.add_paragraph(texte_genere)
        doc.save(filename)

        with open(filename, "rb") as f:
            st.download_button("⬇️ Télécharger le devoir en .docx", f, file_name=os.path.basename(filename))

    except Exception as e:
        st.error(f"Erreur : {e}")